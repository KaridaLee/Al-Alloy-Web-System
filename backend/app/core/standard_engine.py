import json
import sqlite3
from pathlib import Path
from docx import Document
from app.core.config import SQLITE_DIR, BASE_DIR

# 专属企业标准数据库路径
STANDARDS_DB_PATH = SQLITE_DIR / "standards.db"

def init_standards_db():
    """初始化标准数据库，仅保留牌号和成分核心范围"""
    STANDARDS_DB_PATH.parent.mkdir(parents=True, exist_ok=True)
    with sqlite3.connect(STANDARDS_DB_PATH) as conn:
        conn.execute("""
            CREATE TABLE IF NOT EXISTS enterprise_standard (
                brand_name TEXT PRIMARY KEY,
                tech_req_json TEXT,
                ctrl_req_json TEXT,
                updated_at TEXT
            )
        """)

def clean_text(text):
    if text is None:
        return ""
    # 过滤掉换行、空格和特殊的减号、斜杠占位符
    t = str(text).replace("\n", "").replace(" ", "").replace("\r", "").strip()
    return "" if t in ["—", "—", "/", "-"] else t

def sync_word_standard(file_path: str):
    """
    核心 Word 解析引擎：读取第一个成分表
    完美支持从 .docx 结构中百分百稳定提取成分范围
    """
    init_standards_db()
    
    # 从文件名提取牌号 (如从 "QSJ-B02-133 SYZH-AlSi9Cu3企业标准.docx" 提取 "SYZH-AlSi9Cu3")
    filename = Path(file_path).stem
    brand_name = filename.split("企业标准")[0].replace("标准", "").replace("企标", "").strip()
    # 进一步兼容包含编号的前缀
    if " " in brand_name:
        brand_name = brand_name.split(" ")[-1]

    tech_req = {}
    ctrl_req = {}

    try:
        doc = Document(file_path)
    except Exception as e:
        raise ValueError(f"Word文件结构损毁或格式非原生OpenXML: {str(e)}")

    # 扫描文档中出现的第一个表格（即第一页成分表）
    if not doc.tables:
        raise ValueError("该Word文档中没有包含任何受控数据表格")

    table = doc.tables[0]
    
    # 1. 寻找元素表头行
    header_row_idx = -1
    element_cols = {} # 列索引 -> 元素名

    for r_idx, row in enumerate(table.rows):
        row_texts = [clean_text(cell.text) for cell in row.cells]
        # 去重合并，防止合并单元格导致重复计数
        unique_texts = []
        for t in row_texts:
            if t and t not in unique_texts:
                unique_texts.append(t)
        
        found_elements = {}
        for c_idx, cell_text in enumerate(row_texts):
            ct = cell_text.upper()
            el_match = None
            # 标准成分判定字典
            for el in ["SI", "CU", "MG", "MN", "FE", "ZN", "NI", "PB", "SN", "TI", "AL", "CR", "SR", "CA", "SB"]:
                if el in ct:
                    el_match = el.capitalize()
                    break
            if el_match:
                found_elements[c_idx] = el_match

        # 如果某一行包含2个以上的化学元素符号，判定为目标表头行
        if len(found_elements) >= 2:
            header_row_idx = r_idx
            element_cols = found_elements
            break

    if header_row_idx == -1:
        raise ValueError("未能在Word首张表格中检索到包含 Si/Cu/Mg 等化学成分的特征表头")

    # 2. 纵向提取对应列的技术和内控数据范围
    for r_idx in range(header_row_idx + 1, len(table.rows)):
        row = table.rows[r_idx]
        row_texts = [clean_text(cell.text) for cell in row.cells]
        if not row_texts or not row_texts[0]:
            continue

        row_header = row_texts[0]
        if "技术" in row_header or "要求" in row_header:
            for c_idx, el in element_cols.items():
                if c_idx < len(row_texts):
                    val = row_texts[c_idx]
                    if val:
                        tech_req[el] = val
        elif "内控" in row_header or "内部" in row_header:
            for c_idx, el in element_cols.items():
                if c_idx < len(row_texts):
                    val = row_texts[c_idx]
                    if val:
                        ctrl_req[el] = val

    # 持久化存入专属 SQLite 规程库
    from datetime import datetime
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    with sqlite3.connect(STANDARDS_DB_PATH) as conn:
        conn.execute("""
            INSERT INTO enterprise_standard (brand_name, tech_req_json, ctrl_req_json, updated_at)
            VALUES (?, ?, ?, ?)
            ON CONFLICT(brand_name) DO UPDATE SET
                tech_req_json = excluded.tech_req_json,
                ctrl_req_json = excluded.ctrl_req_json,
                updated_at = excluded.updated_at
        """, (
            brand_name,
            json.dumps(tech_req, ensure_ascii=False),
            json.dumps(ctrl_req, ensure_ascii=False),
            now
        ))

def sync_all_standards_in_dir():
    """遍历手动指定的 data/word 目录进行批处理爬取"""
    init_standards_db()
    word_dir = BASE_DIR / "data" / "word"
    if not word_dir.exists():
        word_dir.mkdir(parents=True, exist_ok=True)
        return {"success": False, "message": "已为您创建 data/word 目录，请向其中放入企业标准 Word 文档"}

    # 支持 docx 后缀
    docx_files = list(word_dir.glob("*.docx"))
    # 提示 doc 兼容
    doc_files = list(word_dir.glob("*.doc"))
    
    all_files = docx_files + doc_files
    if not all_files:
        return {"success": False, "message": "data/word 目录中没有检测到任何企业标准文档"}

    success_count = 0
    failed_list = []

    for f_path in all_files:
        try:
            sync_word_standard(str(f_path))
            success_count += 1
        except Exception as e:
            failed_list.append(f"{f_path.name}: {str(e)}")

    return {
        "success": True,
        "total": len(all_files),
        "success_count": success_count,
        "failed_list": failed_list
    }